import json
import sys
import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def parse_html_bold(paragraph, text):
    """
    Parses '<b>...</b>' and applies bold formatting to the runs appropriately.
    """
    parts = re.split(r'(<b>.*?</b>)', text)
    for p in parts:
        if p.startswith('<b>') and p.endswith('</b>'):
            run = paragraph.add_run(p[3:-4])
            run.bold = True
            run.font.name = 'Times New Roman'
        elif p:
            run = paragraph.add_run(p)
            run.font.name = 'Times New Roman'

def set_font(run, size=10, bold=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold

def add_hr(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)

def generate_docx(json_data, output_path):
    document = Document()
    
    # Configure tight margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)

    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(2)
    paragraph_format.space_before = Pt(0)
    paragraph_format.line_spacing = 1.15

    contact = json_data.get('contact', {})
    
    # Header Name
    p_name = document.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_name.add_run(contact.get('name', ''))
    set_font(run_name, size=18, bold=True)
    
    # Contact Info
    p_contact = document.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_parts = []
    
    phone = contact.get('phone')
    email = contact.get('email')
    linkedin = contact.get('linkedin')
    location = contact.get('location')
    
    if phone: contact_parts.append(phone)
    if email: contact_parts.append(email)
    if linkedin: contact_parts.append(linkedin)
    if location: contact_parts.append(location)
    
    p_contact.add_run(" | ".join(contact_parts)).font.size = Pt(10)
    
    def add_section_header(title):
        p = document.add_paragraph()
        run = p.add_run(title.upper())
        set_font(run, size=14, bold=True)
        # Add horizontal line
        add_hr(p)

    # 2. Summary
    if json_data.get('summary'):
        add_section_header('SUMMARY')
        summary_data = json_data['summary']
        if isinstance(summary_data, list):
            for s in summary_data:
                p = document.add_paragraph(style='List Bullet')
                parse_html_bold(p, s)
        elif isinstance(summary_data, str):
            p = document.add_paragraph(style='List Bullet')
            parse_html_bold(p, summary_data)

    # 3. Work Experience
    if json_data.get('experience'):
        add_section_header('WORK EXPERIENCE')
        for exp in json_data['experience']:
            p1 = document.add_paragraph()
            run_title = p1.add_run(exp.get('title', ''))
            set_font(run_title, 11, bold=True)
            
            # Dates
            if exp.get('dates'):
                p1.add_run(f" \t {exp.get('dates')}")
            
            # Line 2: Company - Location
            p2 = document.add_paragraph()
            company = exp.get('company', '')
            loc = exp.get('location', '')
            line2 = f"{company} - {loc}" if loc else company
            run_company = p2.add_run(line2)
            set_font(run_company, 11, bold=False)
            
            # Bullets
            for bullet in exp.get('bullets', []):
                p_bullet = document.add_paragraph(style='List Bullet')
                parse_html_bold(p_bullet, bullet)

    # 4. Other Professional Experience (optional)
    if json_data.get('other_experience'):
        add_section_header('OTHER PROFESSIONAL EXPERIENCE')
        for exp in json_data['other_experience']:
            p1 = document.add_paragraph()
            run_title = p1.add_run(exp.get('title', ''))
            set_font(run_title, 11, bold=True)
            
            # Dates
            if exp.get('dates'):
                p1.add_run(f" \t {exp.get('dates')}")
            
            # Line 2: Company - Location
            p2 = document.add_paragraph()
            company = exp.get('company', '')
            loc = exp.get('location', '')
            line2 = f"{company} - {loc}" if loc else company
            run_company = p2.add_run(line2)
            set_font(run_company, 11, bold=False)
            
            # Bullets
            for bullet in exp.get('bullets', []):
                p_bullet = document.add_paragraph(style='List Bullet')
                parse_html_bold(p_bullet, bullet)

    # 4.5. Projects (optional)
    if json_data.get('projects'):
        add_section_header('PROJECTS')
        for proj in json_data['projects']:
            p = document.add_paragraph()
            run_title = p.add_run(proj.get('name', ''))
            set_font(run_title, 11, bold=True)
            if proj.get('dates'):
                p.add_run(f" | {proj.get('dates')}")
            
            for bullet in proj.get('bullets', []):
                p_bullet = document.add_paragraph(style='List Bullet')
                parse_html_bold(p_bullet, bullet)

    # 5. Education
    if json_data.get('education'):
        add_section_header('EDUCATION')
        for edu in json_data['education']:
            p1 = document.add_paragraph()
            run_deg = p1.add_run(edu.get('degree', ''))
            set_font(run_deg, 11, bold=True)
            if edu.get('dates'):
                p1.add_run(f" \t {edu.get('dates')}")
                
            p2 = document.add_paragraph()
            school = edu.get('school', '')
            loc = edu.get('location', '')
            line2 = f"{school} - {loc}" if loc else school
            p2.add_run(line2)

    # 6. Certifications
    if json_data.get('certifications'):
        add_section_header('CERTIFICATIONS')
        for cert in json_data['certifications']:
            p1 = document.add_paragraph()
            run_cert = p1.add_run(cert.get('name', ''))
            set_font(run_cert, 10, bold=True)
            if cert.get('date'):
                p1.add_run(f" \t {cert.get('date')}")
                
            p2 = document.add_paragraph()
            p2.add_run(cert.get('issuer', ''))

    # 7. Skills
    if json_data.get('skills'):
        add_section_header('SKILLS')
        for skill_cat in json_data['skills']:
            p = document.add_paragraph(style='List Bullet')
            run_cat = p.add_run(f"{skill_cat.get('category', '')}: ")
            set_font(run_cat, 10, bold=True)
            
            # Re-join and parse items
            items_str = ", ".join(skill_cat.get('items', []))
            parse_html_bold(p, items_str)

    document.save(output_path)
    print(f"Successfully generated DOCX: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python generate_resume_docx.py <input.json> <output.docx>")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2]

    if not os.path.exists(input_file):
        print(f"Error: input file '{input_file}' not found.")
        sys.exit(1)

    with open(input_file, 'r', encoding='utf-8') as f:
        data = json.load(f)

    generate_docx(data, output_file)

